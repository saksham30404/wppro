from docx import Document

# Create a new Word document
doc = Document()
doc.add_heading('Machine Failure Prediction Model Guide', level=1)

# Add step-by-step guide sections
sections = [
    ("Step 1: Load the Dataset",
     "Load the dataset using pandas.\n\n"
     "```python\n"
     "import pandas as pd\n\n"
     "# Load the dataset\n"
     "data = pd.read_csv('path_to_your_file/data (1).csv')\n"
     "data.head()\n"
     "```"),

    ("Step 2: Inspect and Understand the Data",
     "It’s crucial to understand the structure and distribution of data before modeling.\n\n"
     "1. **Check Data Types and Missing Values**\n"
     "   ```python\n"
     "   # Summary of dataset\n"
     "   data.info()\n\n"
     "   # Check for missing values\n"
     "   data.isnull().sum()\n"
     "   ```\n\n"
     "2. **Distribution of the Target Variable**\n"
     "   ```python\n"
     "   import seaborn as sns\n"
     "   import matplotlib.pyplot as plt\n\n"
     "   sns.countplot(x='fail', data=data)\n"
     "   plt.title(\"Distribution of Machine Failure (fail)\")\n"
     "   plt.show()\n"
     "   ```"),

    ("Step 3: Exploratory Data Analysis (EDA)",
     "Understanding relationships between features and the target variable is essential.\n\n"
     "1. **Correlation Matrix**\n"
     "   ```python\n"
     "   correlation_matrix = data.corr()\n"
     "   plt.figure(figsize=(10, 8))\n"
     "   sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm')\n"
     "   plt.title(\"Feature Correlation Matrix\")\n"
     "   plt.show()\n"
     "   ```\n\n"
     "2. **Feature Distributions**\n"
     "   ```python\n"
     "   data.hist(bins=20, figsize=(15, 10))\n"
     "   plt.suptitle(\"Feature Distributions\")\n"
     "   plt.show()\n"
     "   ```"),

    ("Step 4: Data Preprocessing",
     "Prepare the data for modeling by splitting it into training and testing sets and scaling if necessary.\n\n"
     "```python\n"
     "from sklearn.model_selection import train_test_split\n"
     "from sklearn.preprocessing import StandardScaler\n\n"
     "# Define features and target\n"
     "X = data.drop(columns='fail')\n"
     "y = data['fail']\n\n"
     "# Split into training and testing sets\n"
     "X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)\n\n"
     "# Standardize the data (if required by your model)\n"
     "scaler = StandardScaler()\n"
     "X_train = scaler.fit_transform(X_train)\n"
     "X_test = scaler.transform(X_test)\n"
     "```"),

    ("Step 5: Train a Machine Learning Model",
     "You can start with a basic logistic regression model, then try other models like Random Forest or Gradient Boosting.\n\n"
     "```python\n"
     "from sklearn.linear_model import LogisticRegression\n"
     "from sklearn.ensemble import RandomForestClassifier\n"
     "from sklearn.metrics import accuracy_score, classification_report\n\n"
     "# Logistic Regression Model\n"
     "lr_model = LogisticRegression()\n"
     "lr_model.fit(X_train, y_train)\n\n"
     "# Predict and evaluate\n"
     "y_pred = lr_model.predict(X_test)\n"
     "print(\"Logistic Regression Accuracy:\", accuracy_score(y_test, y_pred))\n"
     "print(classification_report(y_test, y_pred))\n"
     "```"),

    ("Step 6: Experiment with Different Models",
     "Try models like `RandomForestClassifier` and `GradientBoostingClassifier`, compare their performance, and choose the best.\n\n"
     "```python\n"
     "from sklearn.ensemble import GradientBoostingClassifier\n\n"
     "# Random Forest Model\n"
     "rf_model = RandomForestClassifier(random_state=42)\n"
     "rf_model.fit(X_train, y_train)\n"
     "y_pred_rf = rf_model.predict(X_test)\n"
     "print(\"Random Forest Accuracy:\", accuracy_score(y_test, y_pred_rf))\n\n"
     "# Gradient Boosting Model\n"
     "gb_model = GradientBoostingClassifier(random_state=42)\n"
     "gb_model.fit(X_train, y_train)\n"
     "y_pred_gb = gb_model.predict(X_test)\n"
     "print(\"Gradient Boosting Accuracy:\", accuracy_score(y_test, y_pred_gb))\n"
     "```"),

    ("Step 7: Fine-Tune and Optimize the Model",
     "Use Grid Search or Randomized Search for hyperparameter tuning.\n\n"
     "```python\n"
     "from sklearn.model_selection import GridSearchCV\n\n"
     "# Example: Tune Random Forest parameters\n"
     "param_grid = {\n"
     "    'n_estimators': [50, 100, 200],\n"
     "    'max_depth': [None, 10, 20],\n"
     "    'min_samples_split': [2, 5, 10]\n"
     "}\n\n"
     "grid_search = GridSearchCV(rf_model, param_grid, cv=3, scoring='accuracy')\n"
     "grid_search.fit(X_train, y_train)\n"
     "print(\"Best Parameters for Random Forest:\", grid_search.best_params_)\n"
     "```"),

    ("Step 8: Evaluate the Final Model",
     "After tuning, evaluate the final model on the test data.\n\n"
     "```python\n"
     "best_model = grid_search.best_estimator_\n"
     "y_pred_final = best_model.predict(X_test)\n"
     "print(\"Final Model Accuracy:\", accuracy_score(y_test, y_pred_final))\n"
     "print(classification_report(y_test, y_pred_final))\n"
     "```")
]

# Add sections to the document
for title, content in sections:
    doc.add_heading(title, level=2)
    doc.add_paragraph(content)

# Save the document
doc.save("Machine_Failure_Prediction_Guide.docx")
